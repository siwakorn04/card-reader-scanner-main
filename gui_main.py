import tkinter as tk
from tkinter import messagebox, ttk
from datetime import datetime
from docx import Document
from read_card import SmartCard, APDUCommand
from smartcard.System import readers
import tempfile
import os
import sys
import win32com.client
import re
from docx.shared import Pt
from docx.oxml.ns import qn

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

FORM_FILENAMES = {
    "drive": ("ใบรับรองแพทย์ - ขับขี่", "form_drive.docx"),
    "job": ("ใบรับรองแพทย์ - สมัครงาน", "form_job.docx"),
    "general": ("ใบรับรองแพทย์ - ทั่วไป", "form_general.docx"),
}

def format_thai_date(yyyymmdd):
    try:
        year = int(yyyymmdd[:4])
        month = int(yyyymmdd[4:6])
        day = int(yyyymmdd[6:8])
        thai_months = [
            "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
            "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
        ]
        return f"{day} {thai_months[month - 1]} {year}"
    except:
        return yyyymmdd

def format_cid_boxes(cid):
    return " ".join(list(cid)) if len(cid) == 13 else cid

def fill_medical_certificate(input_docx, output_docx, data):
    doc = Document(input_docx)

    def replace(paragraph):
        text = paragraph.text
        if not any(k in text for k in data):
            return
        paragraph.clear()
        parts = re.split(r"(\{\{.*?\}\})", text)
        for part in parts:
            if part.startswith("{{") and part.endswith("}}"): 
                val = data.get(part.strip(), "")
                run = paragraph.add_run(val)
                run.bold = True
            else:
                run = paragraph.add_run(part)
            run.font.name = "Angsana New"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Angsana New")

    for para in doc.paragraphs:
        replace(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace(para)

    doc.save(output_docx)

def check_reader_status():
    r = readers()
    return "✅ พบเครื่องอ่านบัตร" if r else "❌ ไม่พบเครื่องอ่านบัตร"

def open_form_screen(form_key, branch_name):
    form_name, form_filename = FORM_FILENAMES[form_key]
    form_path = resource_path(f"from/{branch_name}/{form_filename}")

    def read_and_fill():
        try:
            if not all([weight_entry.get(), height_entry.get(), pressure_entry.get(), pulse_entry.get()]):
                messagebox.showwarning("กรอกไม่ครบ", "กรุณากรอก น้ำหนัก ความสูง ความดัน และชีพจร ให้ครบ")
                return

            reader_list = readers()
            if not reader_list:
                messagebox.showerror("ข้อผิดพลาด", "ไม่พบเครื่องอ่านบัตร")
                status_var.set("❌ ไม่พบเครื่องอ่านบัตร")
                status_label.config(fg="red")
                return

            reader = reader_list[0]
            conn = reader.createConnection()
            card = SmartCard(conn)
            card.connect()
            card.initialize()

            commands = {
                "CID": APDUCommand([0x80, 0xB0, 0x00, 0x04, 0x02, 0x00, 0x0D], "CID"),
                "TH Fullname": APDUCommand([0x80, 0xB0, 0x00, 0x11, 0x02, 0x00, 0x64], "TH Fullname"),
                "Date of birth": APDUCommand([0x80, 0xB0, 0x00, 0xD9, 0x02, 0x00, 0x08], "Date of birth"),
                "Address": APDUCommand([0x80, 0xB0, 0x15, 0x79, 0x02, 0x00, 0x64], "Address"),
            }

            card_data = {k: card.read_field(cmd) for k, cmd in commands.items()}

            th_fullname = card_data.get("TH Fullname", "")
            prefix = th_fullname.split()[0] if th_fullname else ""
            cid_raw = card_data.get("CID", "")
            cid_formatted = f"{cid_raw[0]}-{cid_raw[1:5]}-{cid_raw[5:10]}-{cid_raw[10:12]}-{cid_raw[12]}" if len(cid_raw) == 13 else cid_raw
            cid_spaced = format_cid_boxes(cid_raw)
            today = datetime.today()
            ref_id = ref_id_entry.get().strip() or "001"
            
            data = {
                "{{คำนำหน้า}}": prefix,
                "{{ชื่อ}}": th_fullname,
                "{{เลขบัตร}}": cid_raw,
                "{{เลขบัตรช่อง}}": cid_spaced,
                "{{เลขบัตรมีขีด}}": cid_formatted,
                "{{เลขที่}}": ref_id,
                "{{ที่อยู่}}": card_data.get("Address", ""),
                "{{วันเกิด}}": format_thai_date(card_data.get("Date of birth", "")),
                "{{วันที่ตรวจ}}": format_thai_date(today.strftime("%Y%m%d")),
                "{{วัน}}": str(today.day),
                "{{เดือน}}": [
                    "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
                    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
                ][today.month - 1],
                "{{ปี}}": str(today.year + 543),
                "{{หนัก}}": weight_entry.get(),
                "{{สูง}}": height_entry.get(),
                "{{ความดัน}}": pressure_entry.get(),
                "{{ชีพจร}}": pulse_entry.get(),
                "{{สาขา}}": branch_name
            }

            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_docx = os.path.join(tmpdir, "temp.docx")
                tmp_pdf = os.path.join(tmpdir, f"{branch_name}_{ref_id}.pdf")

                fill_medical_certificate(form_path, tmp_docx, data)

                word = win32com.client.Dispatch("Word.Application")
                doc = word.Documents.Open(tmp_docx)
                doc.SaveAs(tmp_pdf, FileFormat=17)
                doc.Close()
                word.Quit()

                if os.path.exists(tmp_pdf):
                    os.startfile(tmp_pdf)
                    messagebox.showinfo("สำเร็จ", "เปิดใบรับรองแพทย์ที่สร้างแล้วเป็น PDF เรียบร้อยแล้ว")
                    form_window.destroy()
                else:
                    messagebox.showerror("ผิดพลาด", "ไม่สามารถสร้าง PDF ได้")

        except Exception as e:
            if '0x8010001D' in str(e):
                messagebox.showerror("ข้อผิดพลาด", "ไม่สามารถเชื่อมต่อกับเครื่องอ่านบัตรได้ กรุณาเปิด Smart Card Service")
            else:
                messagebox.showerror("เกิดข้อผิดพลาด", str(e))

    form_window = tk.Toplevel(root)
    form_window.title(form_name)
    form_window.geometry("450x400")

    tk.Label(form_window, text=f"{form_name} ({branch_name})", font=("TH Sarabun New", 16, "bold")).pack(pady=10)

    frame = tk.Frame(form_window)
    frame.pack(pady=10)

    tk.Label(frame, text="น้ำหนัก (กก.):", font=("TH Sarabun New", 13)).grid(row=0, column=0, sticky="e", padx=5, pady=5)
    weight_entry = tk.Entry(frame)
    weight_entry.grid(row=0, column=1)

    tk.Label(frame, text="ความสูง (ซม.):", font=("TH Sarabun New", 13)).grid(row=1, column=0, sticky="e", padx=5, pady=5)
    height_entry = tk.Entry(frame)
    height_entry.grid(row=1, column=1)

    tk.Label(frame, text="ความดัน (มม.ปรอท):", font=("TH Sarabun New", 13)).grid(row=2, column=0, sticky="e", padx=5, pady=5)
    pressure_entry = tk.Entry(frame)
    pressure_entry.grid(row=2, column=1)

    tk.Label(frame, text="ชีพจร (ครั้ง/นาที):", font=("TH Sarabun New", 13)).grid(row=3, column=0, sticky="e", padx=5, pady=5)
    pulse_entry = tk.Entry(frame)
    pulse_entry.grid(row=3, column=1)

    tk.Label(form_window, text="เลขที่ใบรับรอง:", font=("TH Sarabun New", 13)).pack()
    ref_id_entry = tk.Entry(form_window)
    ref_id_entry.insert(0, "001")
    ref_id_entry.pack()

    tk.Button(form_window, text="✅ อ่านบัตรและกรอกฟอร์ม", font=("TH Sarabun New", 14), command=read_and_fill).pack(pady=20)

    global status_var, status_label
    status_var = tk.StringVar()
    status_var.set(check_reader_status())
    fg_color = "green" if "✅" in status_var.get() else "red"
    status_label = tk.Label(form_window, textvariable=status_var, font=("TH Sarabun New", 12), fg=fg_color)
    status_label.pack(pady=5)

def select_branch():
    def confirm_branch():
        selected = branch_var.get()
        branch_win.destroy()
        show_main(selected)

    branch_win = tk.Tk()
    branch_win.title("เลือกสาขา")
    branch_win.geometry("300x150")

    tk.Label(branch_win, text="กรุณาเลือกสาขา", font=("TH Sarabun New", 16)).pack(pady=10)
    branch_var = tk.StringVar(value="บัวใหญ่")
    ttk.Combobox(branch_win, textvariable=branch_var, values=["บัวใหญ่", "จอหอ"], state="readonly").pack()
    tk.Button(branch_win, text="ยืนยัน", command=confirm_branch).pack(pady=10)

    branch_win.mainloop()

def show_main(branch_name):
    global root
    root = tk.Tk()
    root.title("เลือกประเภทฟอร์มใบรับรองแพทย์")
    root.geometry("400x320")

    tk.Label(root, text=f"คลินิคชีวาดี สาขา{branch_name}", font=("TH Sarabun New", 16)).pack(pady=20)
    tk.Label(root, text="กรุณาเลือกประเภทใบรับรองแพทย์", font=("TH Sarabun New", 16)).pack(pady=20)

    tk.Button(root, text="🚗 ใบรับรองแพทย์ - ขับขี่", font=("TH Sarabun New", 14), width=30,
              command=lambda: open_form_screen("drive", branch_name)).pack(pady=5)
    tk.Button(root, text="💼 ใบรับรองแพทย์ - สมัครงาน", font=("TH Sarabun New", 14), width=30,
              command=lambda: open_form_screen("job", branch_name)).pack(pady=5)

    root.mainloop()

if __name__ == "__main__":
    select_branch()
