from read_card import SmartCard, APDUCommand, select_reader, thai2unicode
from docx import Document
from datetime import datetime
import sys

# ✅ ฟังก์ชันกรอกข้อมูลลงในฟอร์ม .docx (รองรับ placeholder ข้าม run)
def fill_medical_certificate(input_docx, output_docx, data):
    doc = Document(input_docx)

    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for key, val in data.items():
            if key in full_text:
                new_text = full_text.replace(key, val)
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

    # แทนที่ในย่อหน้า
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # แทนที่ในตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(output_docx)
    print("✅ สร้างใบรับรองแพทย์แล้ว:", output_docx)

# 🔁 แปลงวันที่จาก YYYYMMDD → 19 กุมภาพันธ์ 2547
def format_thai_date(yyyymmdd):
    try:
        year = int(yyyymmdd[:4])
        month = int(yyyymmdd[4:6])
        day = int(yyyymmdd[6:8])
        thai_months = [
            "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
            "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
        ]
        return f"{day} {thai_months[month - 1]} {year}"
    except:
        return yyyymmdd

# 🚀 ฟังก์ชันหลัก
def main():
    try:
        reader = select_reader()
        if reader is None:
            sys.exit(1)

        conn = reader.createConnection()
        card = SmartCard(conn)
        card.connect()
        card.initialize()

        # คำสั่งดึงข้อมูลจากบัตร
        commands = {
            "CID": APDUCommand([0x80, 0xB0, 0x00, 0x04, 0x02, 0x00, 0x0D], "CID"),
            "TH Fullname": APDUCommand([0x80, 0xB0, 0x00, 0x11, 0x02, 0x00, 0x64], "TH Fullname"),
            "Date of birth": APDUCommand([0x80, 0xB0, 0x00, 0xD9, 0x02, 0x00, 0x08], "Date of birth"),
            "Address": APDUCommand([0x80, 0xB0, 0x15, 0x79, 0x02, 0x00, 0x64], "Address"),
        }

        card_data = {}
        for key, cmd in commands.items():
            card_data[key] = card.read_field(cmd)

        # แปลงวัน
        today_thai = format_thai_date(datetime.today().strftime("%Y%m%d"))
        dob = format_thai_date(card_data.get("Date of birth", ""))

        # เตรียมข้อมูลสำหรับกรอกลงฟอร์ม
        data = {
            "{{ชื่อ}}": card_data.get("TH Fullname", ""),
            "{{เลขบัตร}}": card_data.get("CID", ""),
            "{{ที่อยู่}}": card_data.get("Address", ""),
            "{{วันเกิด}}": dob,
            "{{วันที่ตรวจ}}": today_thai,
            "{{วันที่ออก}}": today_thai
        }

        input_doc = "ใบรับรองแพทยฺ์(ใยอนุญาตขับรถ) บัวใหญ่ test.docx"
        output_doc = "ใบรับรองแพทย์_กรอกแล้ว.docx"
        fill_medical_certificate(input_doc, output_doc, data)

    except Exception as e:
        print("❌ เกิดข้อผิดพลาด:", e)

if __name__ == "__main__":
    main()
