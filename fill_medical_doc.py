from docx import Document

def fill_medical_certificate(input_docx, output_docx, data):
    doc = Document(input_docx)

    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for key, val in data.items():
            if key in full_text:
                new_text = full_text.replace(key, val)
                # ล้างทุก run ก่อน แล้วใส่ข้อความใหม่ใน run แรก
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

    # แทนที่ในย่อหน้า
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # แทนที่ในตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(output_docx)
    print(f"✅ กรอกฟอร์มสำเร็จ: {output_docx}")
